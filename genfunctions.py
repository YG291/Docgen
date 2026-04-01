from docx import Document
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.table import Table
from pathlib import Path
from typing import Any

from doc_gen import Doc
from doc_gen import template_path
from conversion_script import fix_alias, load_data_from_csv

def directors_resolutions(global_dict: dict, people: dict):
    path = template_path("Directors_ Resolutions", "initial directors resolution.docx")
    doc = Doc(path, global_dict, people)
    doc.replace_main(global_dict)
    sale_table = doc.find_table_by_text('[BUYER]')
    for person in people['shareholders']:
        new_row = doc.copy_row(sale_table, 1)
        for cell in new_row.cells:
            for p in cell.paragraphs:
                doc._replace_p(p, people['shareholders'][person], '[BUYER]')
                doc._replace_p(p, people['shareholders'][person], '[COUNT]')
                doc._replace_p(p, people['shareholders'][person], '[PRICE]')
    doc.remove_row(sale_table, 1)
    doc.fill_signatures(people['directors'])
    doc.save_doc("Directors_ Resolutions", "initial directors resolution.docx", global_dict['[CORP]'])

def indemnification_agreement(global_dict: dict, people: dict):
    for person in people['directors']:
        path = template_path("Indemnification", "indemnification agreement.docx")
        doc = Doc(path, global_dict, people)
        doc.replace_main(global_dict)
        doc.replace_tables(global_dict)
        doc.replace_tables(people['directors'][person])
        doc.replace_main(people['directors'][person])
        doc.save_doc("Indemnification", "indemnification "+person+" .docx", global_dict['[CORP]'])
    
    for person in people['officers']:
        path = template_path("Indemnification", "indemnification agreement.docx")
        doc = Doc(path, global_dict, people)
        doc.replace_main(global_dict)
        doc.replace_main(people['officers'][person])
        doc.replace_tables(global_dict)
        doc.replace_tables(people['officers'][person])
        doc.save_doc("Indemnification", "indemnification "+person+" .docx", global_dict['[CORP]'])

def bylaws(global_dict: dict, people: dict):
    path = template_path("By-Laws", "bylaws.docx")
    doc = Doc(path, global_dict, people)
    doc.replace_main(global_dict)
    doc.save_doc("By-Laws", "bylaws.docx", global_dict['[CORP]'])

    path = template_path("By-Laws", "certificate of secretary.docx")
    doc = Doc(path, global_dict, people)
    doc.replace_main(global_dict)
    doc.replace_tables(global_dict)
    doc.save_doc("By-Laws", "certificate of secretary.docx", global_dict['[CORP]'])

    path = template_path("By-Laws", "incorporation certificate.docx")
    doc = Doc(path, global_dict, people)
    doc.replace_main(global_dict)
    doc.save_doc("By-Laws", "incorporation certificate.docx", global_dict['[CORP]'])

def ipassignment(global_dict: dict, people: dict):
    path = template_path("IP Assignments", "ip assignment.docx")
    for person in people['ip']:
        doc = Doc(path, global_dict, people)
        doc.replace_hf(global_dict)
        doc.replace_hf(people['ip'][person])
        doc.replace_main(global_dict)
        doc.replace_main(people['ip'][person])
        doc.replace_tables(people['ip'][person])
        doc.replace_tables(global_dict)
        p = doc.doc.paragraphs[len(doc.doc.paragraphs)-1]
        doc.dot_jots(p, people['ip'][person]['item_list'])
        doc.save_doc("IP Assignments", "ip assignment "+person+" .docx", global_dict['[CORP]'])

def jointescrow(global_dict: dict, people: dict):
     path = template_path("Joint Escrow Instructions", "joint escrow instructions.docx")
     for person in people['shareholders']:
        doc = Doc(path, global_dict, people)
        doc.replace_hf(global_dict)
        doc.replace_hf(people['shareholders'][person])
        doc.replace_main(global_dict)
        doc.replace_main(people['shareholders'][person])
        doc.replace_tables(people['shareholders'][person])
        doc.replace_tables(global_dict)
        doc.save_doc("Joint Escrow Instructions", "joint escrow "+person+" .docx", global_dict['[CORP]'])

def restrictedpurchase(global_dict: dict, people: dict):
    path = template_path("Restricted Stock Purchase", "restricted stock purchase agreement.docx")
    for person in people['shareholders']:
        doc = Doc(path, global_dict, people)
        doc.replace_hf(global_dict)
        doc.replace_hf(people['shareholders'][person])
        doc.replace_main(global_dict)
        doc.replace_main(people['shareholders'][person])
        doc.replace_tables(people['shareholders'][person])
        doc.replace_tables(global_dict)
        doc.save_doc("Restricted Stock Purchase", "restricted stock purchase "+person+" .docx", global_dict['[CORP]'])

def stock_assignment(global_dict: dict, people: dict):
    path = template_path("Restricted Stock Purchase", "stock assignment.docx")
    for person in people['shareholders']:
        doc = Doc(path, global_dict, people)
        doc.replace_main(global_dict)
        doc.replace_main(people['shareholders'][person])
        doc.replace_tables(people['shareholders'][person])
        doc.replace_tables(global_dict)
        doc.save_doc("Restricted Stock Purchase", "stock assignment "+person+" .docx", global_dict['[CORP]'])

import copy
from docx.table import Table

def shareholder_resolution(global_dict: dict, people: dict):
    path = template_path("Shareholders_ Resolutions", "shareholders resolutions.docx")
    doc = Doc(path, global_dict, people)
    
    # 1. Replace headers, footers, and main body text globally
    doc.replace_hf(global_dict)
    doc.replace_main(global_dict)

    # 2. Find the template table
    template_table = doc.find_table_by_text('[BUYER]')
    
    if template_table:
        for person_name in people['shareholders']:
            # THE CLONE: Copy the entire table XML
            new_xml = copy.deepcopy(template_table._element)
            template_table._element.addnext(new_xml)
            new_table = Table(new_xml, doc.doc)
            
            # 3. Data Mapping: Merge person info with global info for this block
            # This ensures we have [BUYER] AND [INCDATE] available for every table
            p_data = people['shareholders'][person_name].copy()
            p_data['[BUYER]'] = person_name
            p_data['[SHAREHOLDER]'] = person_name
            p_data['[STOCKHOLDER]'] = person_name
            
            # 4. SURGICAL REPLACEMENT: Cell by Cell
            for row in new_table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        # Check for Person-Specific Tags
                        for tag in ['[BUYER]', '[SHAREHOLDER]', '[STOCKHOLDER]', '[COUNT]', '[PRICE]']:
                            if tag in p.text and tag in p_data:
                                doc._replace_p(p, p_data, tag)
                        
                        # Check for Global/Corporate Tags (like [INCDATE], [DATE], [CORP])
                        for g_tag in global_dict:
                            if g_tag in p.text:
                                doc._replace_p(p, global_dict, g_tag)

        # 5. CLEANUP: Kill the placeholder table after N copies are made
        doc._remove_table(template_table)
        
    else:
        # Emergency fallback
        doc.fill_signatures(people['shareholders'])

    # 6. Save using the Corporate Name
    doc.save_doc("Shareholders_ Resolutions", "shareholders resolutions.docx", global_dict['[CORP]'])

def stock_certificates(global_dict: dict, people: dict):
    counter = 1
    for person in people['shareholders']:
        path = template_path("Stock Certificates", "stock certificates.docx")
        doc = Doc(path, global_dict, people)
        doc.replace_main(global_dict)
        doc.replace_main(people['shareholders'][person])
        doc.replace_tables(people['shareholders'][person])
        doc.replace_tables(global_dict)
        doc.replace_tables({'[CERTNUM]': str(counter)})
        doc.save_doc("Stock Certificates", "stock certificate "+person+" .docx", global_dict['[CORP]'])
        counter += 1

def stock_purchase_agreement(global_dict: dict, people: dict):
    for person in people['shareholders']:
        path = template_path("Stock Purchase Agreement", "stock purchase agreement.docx")
        doc = Doc(path, global_dict, people)
        doc.replace_main(global_dict)
        doc.replace_main(people['shareholders'][person])
        doc.replace_tables(people['shareholders'][person])
        doc.replace_tables(global_dict)
        doc.save_doc("Stock Purchase Agreement", "stock purchase agreement "+person+" .docx", global_dict['[CORP]'])


if __name__ == '__main__':
    global_csv_path = Path(__file__).parent / "input" / "global_data.csv"
    people_csv_path = Path(__file__).parent / "input" / "people.csv"
    global_dict, people = load_data_from_csv(global_csv_path, people_csv_path)

    directors_resolutions(global_dict, people)
    indemnification_agreement(global_dict, people)
    bylaws(global_dict, people)
    ipassignment(global_dict, people)
    jointescrow(global_dict, people)
    restrictedpurchase(global_dict, people)
    stock_assignment(global_dict, people)
    shareholder_resolution(global_dict, people)
    stock_certificates(global_dict, people)
    stock_purchase_agreement(global_dict, people)