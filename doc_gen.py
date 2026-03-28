from docx import Document
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.table import Table
from pathlib import Path

from typing import Any

class Doc:
    """Representation of a generic document.
    """
    
    def __init__(self, name: str, globaldict: dict, peopledict: dict, settingsdict={'cert_start': 1} ) -> None:
        #people comes in the form of {'directors': {sam: {name, address, etc}}, 'shareholders': {eric: name, address, etc}, 'officers': {jacob: name, address, etc}, 'ip': {jackson: name, address, etc}}
        self.doc = Document(name)
        self.globaldict = globaldict
        self.people = peopledict
        self.settings = settingsdict

    def _replace_in_run(self, r: Run, old: str, new: str) -> None:
        """Helper method that replaces every instance of old in r.text with new.
        """
        r.text = r.text.replace(old, new)

    def _replace_p(self, p: Paragraph, values: dict, key: str) -> None:
        """Helper method that fixes every fragmented run in paragraph <<p>> and replaces every instance 
        of key with its corresponding value in <<values>>.
        """
        while key in p.text:
            start_index = p.text.find(key)
            end_index = start_index + len(key)
            current = 0
            _runs = []
            for r in p.runs:
                if current < end_index and current + len(r.text) >= start_index:
                    _runs.append(r)
                current += len(r.text)
                if current > end_index:
                    break
            for r_index in range(len(_runs)):
                if r_index != 0: 
                    _runs[0].text = _runs[0].text + _runs[r_index].text
                    _runs[r_index].text = ''
            if _runs:
               self._replace_in_run(_runs[0], key, values[key])

    def replace_main(self, values: dict) -> None:
        """Replaces every instance of keys from <<values>> in the document's primary paragraphs.
        """
        for key in values:
            for p in self.doc.paragraphs:
                self._replace_p(p, values, key)

    def _replace_t(self, t: Table, values: dict, key) -> None:
        """Helper method to replace paragraphs in tables.
        """
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    self._replace_p(p, values, key)

    def replace_tables(self, values: dict) -> None:
        """Replaces every instance of keys from <<values>> in the document's tables.
        """
        for key in values:
            for t in self.doc.tables:
                self._replace_t(t, values, key)
        
    def replace_hf(self, values: dict) -> None:
        """Replaces every instance of keys from <<values>> in the document's headers and footers.
        """
        for s in self.doc.sections:
            for p in s.footer.paragraphs:
                for key in values:
                    self._replace_p(p, values, key)
            for t in s.footer.tables:
                for key in values:    
                    self._replace_t(t, values, key)
            for p in s.first_page_footer.paragraphs:
                for key in values:
                    self._replace_p(p, values, key)
            for t in s.first_page_footer.tables:
                for key in values:
                    self._replace_t(t, values, key)
            for p in s.header.paragraphs:
                for key in values:
                    self._replace_p(p, values, key)
            for t in s.header.tables:
                for key in values:
                    self._replace_t(t, values, key)

    def _find_last_table(self) -> Table | None:
        if self.doc.tables is not None and self.doc.tables != []:
            return self.doc.tables[len(self.doc.tables) - 1]
        return None
        
    def copy_table(self, tb: Table) -> Table:
        """Adds a table immediately after tb._element with merged-cell support."""
        colnum, rownum = len(tb.columns), len(tb.rows)
        new_tb = self.doc.add_table(rownum, colnum)
        tb._element.addnext(new_tb._element)
        
        # Use zip to iterate through actual existing cells in both tables simultaneously
        for old_row, new_row in zip(tb.rows, new_tb.rows):
            for old_cell, new_cell in zip(old_row.cells, new_row.cells):
                # 1. Align paragraph count (skip the first one created by default)
                for i in range(len(old_cell.paragraphs) - 1):
                    new_cell.add_paragraph()
                
                # 2. Copy the content and formatting
                for p_index, old_p in enumerate(old_cell.paragraphs):
                    new_p = new_cell.paragraphs[p_index]
                    for r in old_p.runs:
                        new_run = new_p.add_run(r.text)
                        # Copy styles/formatting
                        new_run.style = r.style
                        new_run.bold = r.bold
                        new_run.italic = r.italic
                        if r.font.name: new_run.font.name = r.font.name
                        if r.font.size: new_run.font.size = r.font.size
        return new_tb
    
    def _remove_table(self, table: Table) -> None:
        tb = table._element
        tb.getparent().remove(tb)
    
    def fill_signatures(self, who: dict) -> None:
        tb = self._find_last_table()
        index = self.doc.tables[len(self.doc.tables) - 1]
        for person in who.values():
            table = self.copy_table(tb)
            for key in person:
                self._replace_t(table, person, key)
        self._remove_table(tb)

    
    def find_table_by_text(self, txt: str) -> Table | None:
        """Returns the first table containing the trigger_text."""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if txt in cell.text:
                        return table
                    
    def remove_row(self, table: Table, row_index: int) -> None:
        """Removes a row from a table at the specified index."""
        tbl = table._element
        tr = table.rows[row_index]._element
        tbl.remove(tr)
    
    def copy_row(self, table: Table, index: int = 1) -> Any:
        """Adds a new row and copies formatting from the template row."""
        new_row = table.add_row()
        template_row = table.rows[index]
        
        for i, cell in enumerate(new_row.cells):
            source_cell = template_row.cells[i]
            # Clear default paragraph created by add_row()
            cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
            
            for source_p in source_cell.paragraphs:
                target_p = cell.add_paragraph()
                target_p.style = source_p.style
                target_p.alignment = source_p.alignment
                # COPY THE RUNS so the placeholders actually exist in the new row
                for source_run in source_p.runs:
                    new_run = target_p.add_run(source_run.text)
                    new_run.bold = source_run.bold
                    new_run.italic = source_run.italic
                    new_run.font.name = source_run.font.name
                    new_run.font.size = source_run.font.size
        return new_row
        

    def _delete_paragraph(self, p: Paragraph) -> None:
        para = p._element
        para.getparent().remove(para)

    def dot_jots(self, p: Paragraph, item_list: list) -> None:
        nparalist = []
        for i in item_list:
            new_p = self.doc.add_paragraph('', p.style)
            p._element.addnext(new_p._element)
            nparalist.append(new_p)
        nparalist.reverse()
        for j in range(len(item_list)):
            nparalist[j].add_run(str(item_list[j]))
        self._delete_paragraph(p)

    def save_doc(self, folder, savename) -> None:
        main_path = Path(__file__).parent
        output_dir = main_path / "output" / folder
        output_dir.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_dir / savename)
    
def template_path(folder: str, file: str) -> Path:
    main_path = Path(__file__).parent
    doc_path = main_path / "templates" / folder / file
    return doc_path